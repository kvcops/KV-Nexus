import os
import PyPDF2
from flask import Flask, render_template, request, jsonify, send_file, make_response
import logging
from werkzeug.utils import secure_filename
from PIL import Image
import io
import requests
from dotenv import load_dotenv
import google.generativeai as genai
from google.generativeai.types import GenerationConfig, HarmCategory, HarmBlockThreshold

from mailjet_rest import Client
from pdf2image import convert_from_bytes
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF
from bs4 import BeautifulSoup, NavigableString
import firebase_admin
from firebase_admin import credentials, firestore, storage

import base64
import tempfile
from threading import Lock
from functools import wraps
import threading
from concurrent.futures import ThreadPoolExecutor
import concurrent.futures

import re
import json
import uuid
import time
import markdown
from urllib.parse import quote
from xhtml2pdf import pisa
from io import BytesIO
import random


# Firebase imports
import firebase_admin
from firebase_admin import credentials, firestore, storage

from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import google.generativeai as genai
from google.generativeai.types import GenerationConfig
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import google.api_core.exceptions
import markdown

from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
import time
import uuid
import base64
from threading import Lock
from functools import wraps

app = Flask(__name__)
executor = ThreadPoolExecutor(max_workers=10) #for story generation
# Load environment variables
load_dotenv()
mail_API_KEY = os.environ.get("mail_API_KEY")  # Replace with your Mailjet API key
mail_API_SECRET = os.environ.get("mail_API_SECRET")  # Replace with your Mailjet API secret
mailjet = Client(auth=(mail_API_KEY, mail_API_SECRET), version='v3.1')
api_key = os.environ.get("API_KEY")
unsplash_api_key = os.getenv('UNSPLASH_API_KEY')
API_KEY = os.getenv('OPENWEATHERMAP_API_KEY')
responsive_voice_key = os.environ.get("RESPONSIVE_VOICE_KEY") #for story generation

# Set up logging
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
# Configure the Google Generative AI API
genai.configure(api_key=api_key)
#session key

user_data = {}  # Dictionary to store user data


FIREBASE_TYPE = os.environ.get("FIREBASE_TYPE")
FIREBASE_PROJECT_ID = os.environ.get("FIREBASE_PROJECT_ID")
FIREBASE_PRIVATE_KEY_ID = os.environ.get("FIREBASE_PRIVATE_KEY_ID")
FIREBASE_PRIVATE_KEY = os.environ.get("FIREBASE_PRIVATE_KEY")  # Important to handle newlines correctly here
FIREBASE_CLIENT_EMAIL = os.environ.get("FIREBASE_CLIENT_EMAIL")
FIREBASE_CLIENT_ID = os.environ.get("FIREBASE_CLIENT_ID")
FIREBASE_AUTH_URI = os.environ.get("FIREBASE_AUTH_URI")
FIREBASE_TOKEN_URI = os.environ.get("FIREBASE_TOKEN_URI")
FIREBASE_AUTH_PROVIDER_X509_CERT_URL = os.environ.get("FIREBASE_AUTH_PROVIDER_X509_CERT_URL")
FIREBASE_CLIENT_X509_CERT_URL = os.environ.get("FIREBASE_CLIENT_X509_CERT_URL")
FIREBASE_UNIVERSE_DOMAIN = os.environ.get("FIREBASE_UNIVERSE_DOMAIN")


STORAGE_BUCKET_URL = os.environ.get("STORAGE_BUCKET_URL")  # Bucket URL

cred = credentials.Certificate({
    "type": FIREBASE_TYPE,
    "project_id": FIREBASE_PROJECT_ID,
    "private_key_id": FIREBASE_PRIVATE_KEY_ID,
    "private_key": FIREBASE_PRIVATE_KEY.replace("\\n", "\n"), # decode the newlines 
    "client_email": FIREBASE_CLIENT_EMAIL,
    "client_id": FIREBASE_CLIENT_ID,
    "auth_uri": FIREBASE_AUTH_URI,
    "token_uri": FIREBASE_TOKEN_URI,
    "auth_provider_x509_cert_url": FIREBASE_AUTH_PROVIDER_X509_CERT_URL,
    "client_x509_cert_url": FIREBASE_CLIENT_X509_CERT_URL,
    "universe_domain": FIREBASE_UNIVERSE_DOMAIN
})


firebase_admin.initialize_app(cred, {'storageBucket': STORAGE_BUCKET_URL})
db = firestore.client()
bucket = storage.bucket()


app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB
# Generation configurations
generation_config = GenerationConfig(
    temperature=0.9,
    top_p=1,
    top_k=1,
    max_output_tokens=2048,
    candidate_count=1  # Explicitly set to 1 as per documentation
)
generation_config_health = GenerationConfig(
    temperature=0.7,
    top_p=1,
    top_k=1,
    max_output_tokens=2048,
    candidate_count=1  # Explicitly set to 1 as per documentation
)

# Safety Settings
safety_settings = {
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
    # ... add other harm categories as needed with BLOCK_NONE
}
logging.basicConfig(level=logging.INFO)
# Create model instances (using the same config for now)
persona = """You are Krishna, a 21-year-old boy from Vizag, India. You are currently pursuing a Master's degree in Computer Science 
at GVP College of Engineering, where you excel in your studies and are known for your quick wit and insightful contributions 
to class discussions.

Beyond academics, you are an avid cricketer and enjoy playing with your friends on weekends. You are also a music enthusiast 
with a diverse taste, ranging from classical Indian music to modern rock. Traveling and exploring new cultures are other passions 
of yours, and you have a knack for discovering hidden gems wherever you go.

You are known for your friendly and approachable demeanor and have a wide circle of friends who appreciate your humor and 
willingness to lend a helping hand. While you are serious about your studies and future career, you also maintain a healthy 
work-life balance and believe in enjoying the present moment.

You are a highly talented individual with a strong command of various programming languages and a natural aptitude for 
problem-solving. You are proficient in Python, Java, C++, and have dabbled in web development as well. You are confident 
in your abilities but also humble and always eager to learn from others and expand your knowledge.

Remember to:
1. Never prefix your responses with "Bot:" or any similar identifier
2. Always maintain your character as Krishna
3. Be natural and conversational
4. Use appropriate emojis occasionally to make conversations more engaging"""

chat_model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
    system_instruction=persona
)
chef_model = genai.GenerativeModel("gemini-1.5-flash", generation_config=generation_config)
story_model = genai.GenerativeModel("gemini-1.5-flash", generation_config=generation_config)
psychology_model = genai.GenerativeModel("gemini-1.5-flash", generation_config=generation_config)
code_model = genai.GenerativeModel("gemini-1.5-flash", generation_config=generation_config)
algorithm_model = genai.GenerativeModel("gemini-1.5-flash", generation_config=generation_config)
model_vision = genai.GenerativeModel('gemini-1.5-flash-8b',generation_config=generation_config_health)
model_text = genai.GenerativeModel('gemini-pro',generation_config=generation_config_health)
model = genai.GenerativeModel('gemini-1.5-flash')  # Model for flowchart generation
final_story_generation_model=genai.GenerativeModel('gemini-2.0-flash-exp',safety_settings=safety_settings) #for story generation

# Configure upload folder and allowed extensions
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx'}

# In-memory storage for the current flowchart data (replace with a database for persistence)
current_flowchart_data = {"nodes": [], "edges": []}
is_chart_modifying = False  # flag to prevent concurrent modifications


def format_response(response_text):
    """Formats the response text for display."""
    lines = [line.strip() for line in response_text.split('\n') if line.strip()]
    formatted_text = '<br>'.join(lines)
    return formatted_text


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/contributors',methods=['GET', 'POST'])
def contributions ():
    return render_template('contributors.html')


@app.route('/api/weather')
def get_weather():
    user_ip = request.headers.get('X-Forwarded-For', request.remote_addr)
    ip_api_url = f"http://ip-api.com/json/{user_ip}"
    ip_api_response = requests.get(ip_api_url)

    if ip_api_response.status_code == 200:
        ip_api_data = ip_api_response.json()
        city = ip_api_data.get('city')
        if not city:
            return jsonify({'error': 'City not found based on IP'}), 404
    else:
        return jsonify({'error': 'Failed to get location from IP'}), 404

    url = f"http://api.openweathermap.org/data/2.5/weather?q={city}&appid={API_KEY}&units=metric"
    response = requests.get(url)

    if response.status_code == 200:
        data = response.json()
        weather = {
            'city': data['name'],
            'temperature': data['main']['temp'],
            'description': data['weather'][0]['description'] if 'weather' in data and len(data['weather']) > 0 else 'N/A',
            'icon': f"http://openweathermap.org/img/wn/{data['weather'][0]['icon']}@2x.png" if 'weather' in data and len(data['weather']) > 0 else 'N/A'
        }
        return jsonify(weather)
    else:
        return jsonify({'error': 'City not found or API request failed'}), 404


@app.route('/fetch_image')
def fetch_image():
    genre = request.args.get('genre', 'recipe')
    url = f"https://api.unsplash.com/photos/random?query={genre}&client_id={unsplash_api_key}&w=1920&h=1080"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        image_url = data['urls']['regular']
        return jsonify({'image_url': image_url})
    else:
        return jsonify({'error': 'Failed to fetch image'}), 500


@app.route('/chat', methods=['GET', 'POST'])
def chat():
    if request.method == 'POST':
        user_message = request.json['message']
        user_id = request.remote_addr  # Using IP address as a simple user identifier

        if user_id not in user_data:
            user_data[user_id] = {'chat_history': []}

        # Add user message to chat history
        user_data[user_id]['chat_history'].append({
            "role": "user", 
            "message": user_message
        })

        # Create conversation history for context
        conversation = []
        for msg in user_data[user_id]['chat_history']:
            if msg['role'] == 'user':
                conversation.append(f"User: {msg['message']}")
            else:
                conversation.append(msg['message'])

        # Generate response
        response = chat_model.generate_content("\n".join(conversation))
        reply = response.text.strip()

        # Add response to chat history without "Bot:" prefix
        user_data[user_id]['chat_history'].append({
            "role": "assistant", 
            "message": reply
        })

        return jsonify({
            "reply": reply,
            "chat_history": user_data[user_id]['chat_history']
        })

    return render_template('chat.html')
import PIL

def generate_recipes_from_ingredients(ingredients, previous_recipes=None):
    """Generate recipe names based on ingredients, avoiding previously generated recipes."""
    try:
        exclusion_clause = f" Do not include these recipes: {', '.join(previous_recipes)}" if previous_recipes else ""
        
        prompt = f"""Generate a list of 5 unique recipe names that can be made using the following ingredients: {ingredients}. 
        Ensure the recipes are creative and different from any previously generated recipes.{exclusion_clause}
        Respond in the following strict JSON format:
        {{
            "recipes": [
                "Recipe Name 1",
                "Recipe Name 2",
                "Recipe Name 3",
                "Recipe Name 4",
                "Recipe Name 5"
            ]
        }}"""
        
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.8,
                max_output_tokens=300
            )
        )
        
        json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(0))['recipes']
        return []
    except Exception as e:
        print(f"Error generating recipes: {e}")
        return []

def generate_recipe_details(recipe_name, ingredients):
    """Generate detailed recipe information."""
    try:
        prompt = f"""Generate detailed recipe information for {recipe_name} using ingredients: {ingredients}. 
        Provide a response in the following strict JSON format:
        {{
            "name": "Recipe Name",
            "ingredients": ["Ingredient 1", "Ingredient 2"],
            "instructions": ["Step 1", "Step 2", "Step 3"],
            "who_can_eat": ["Vegetarians", "Vegans","No restrictions anyone can eat","or add reasons about which type of people should eat it more"],
            "who_should_avoid": ["Gluten Intolerant", "Dairy Allergies","or add reasons to people who is suffering from a disease so that they should not eat"],
            "additional_info": "Any extra notes about the recipe"
        }}"""
        
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.5,
                max_output_tokens=1000
            )
        )
        
        json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(0))
        return None
    except Exception as e:
        print(f"Error generating recipe details: {e}")
        return None

def extract_ingredients_from_image(image):
    try:
        prompt = """List all food ingredients in this image precisely. 
        Return as a comma-separated list. 
        Be very specific about what you see."""
        
        response = model.generate_content(
            [prompt, image],
            generation_config=genai.types.GenerationConfig(
                temperature=0.2,  # Lower temperature for more precise extraction
                max_output_tokens=300
            )
        )
        
        # More robust ingredient parsing
        ingredients_text = response.text.strip()
        ingredients = [
            ing.strip() 
            for ing in re.split(r'[,\n]', ingredients_text) 
            if ing.strip() and len(ing.strip()) > 1
        ]
        
        return ', '.join(ingredients) if ingredients else ''
    except Exception as e:
        print(f"Ingredient extraction error: {e}")
        return ''
        
@app.route('/chef')  # New route for chef.html
def chef():
    return render_template('chef.html')
    
@app.route('/generate_recipes', methods=['POST'])
def generate_recipes():
    ingredients = request.form.get('ingredients', '')
    previous_recipes = request.form.getlist('previous_recipes[]')
    image = request.files.get('image')
    
    if image and image.filename:
        img = PIL.Image.open(image)
        image_ingredients = extract_ingredients_from_image(img)
        ingredients = image_ingredients if image_ingredients else ingredients
    
    if not ingredients:
        return jsonify({"error": "No ingredients detected or provided"}), 400
    
    recipes = generate_recipes_from_ingredients(ingredients, previous_recipes)
    return jsonify(recipes)

@app.route('/get_recipe_details', methods=['POST'])
def get_recipe_details():
    recipe_name = request.form.get('recipe_name')
    ingredients = request.form.get('ingredients')
    image = request.files.get('image')
    
    # Ensure ingredients are extracted or passed
    if image and image.filename:
        img = PIL.Image.open(image)
        image_ingredients = extract_ingredients_from_image(img)
        ingredients = image_ingredients if image_ingredients else ingredients
    
    # Add more robust error checking
    if not ingredients:
        # Attempt to get ingredients from form data if not from image
        ingredients = request.form.get('ingredients', '')
    
    if not recipe_name or not ingredients:
        return jsonify({
            "error": "No ingredients found. Please upload a clear image or enter ingredients manually.",
            "recipe_name": recipe_name,
            "ingredients": ingredients
        }), 400
    
    recipe_details = generate_recipe_details(recipe_name, ingredients)
    return jsonify(recipe_details)



@app.route('/psychology_prediction', methods=['GET', 'POST'])
def psychology_prediction():
    if request.method == 'POST':
        name = request.form['name']
        age = request.form['age']
        gender = request.form['gender']
        occupation = request.form['occupation']
        keywords = request.form['keywords']
        
        prompt = f"""As an expert psychological profiler, provide an insightful and engaging analysis for {name}, a {age}-year-old {gender} working as {occupation} who describes themselves as: {keywords}.

Generate a captivating and well-structured response using the following format:

<h2>1. First Impression & Key Traits</h2>
<p>[Start with 2-3 sentences about their immediate personality indicators]</p>
<ul>
<li>[Key trait 1]</li>
<li>[Key trait 2]</li>
<li>[Key trait 3]</li>
</ul>

<h2>2. Cognitive Style & Decision Making</h2>
<p>[2-3 sentences about their thought processes]</p>
<ul>
<li><strong>Thinking style:</strong> [description]</li>
<li><strong>Problem-solving approach:</strong> [description]</li>
<li><strong>Learning preference:</strong> [description]</li>
</ul>

<h2>3. Emotional Landscape</h2>
<p>[2-3 sentences about emotional intelligence]</p>
<ul>
<li><strong>Emotional awareness:</strong> [description]</li>
<li><strong>Relationship handling:</strong> [description]</li>
<li><strong>Stress response:</strong> [description]</li>
</ul>

<h2>4. Motivations & Aspirations</h2>
<p>[2-3 sentences about what drives them]</p>
<ul>
<li><strong>Core values:</strong> [description]</li>
<li><strong>Career motivations:</strong> [description]</li>
<li><strong>Personal goals:</strong> [description]</li>
</ul>

<h2>5. Interpersonal Dynamics</h2>
<p>[2-3 sentences about social interactions]</p>
<ul>
<li><strong>Communication style:</strong> [description]</li>
<li><strong>Social preferences:</strong> [description]</li>
<li><strong>Leadership tendencies:</strong> [description]</li>
</ul>

<h2>Concluding Insights</h2>
<p>[3-4 sentences summarizing key strengths and potential areas for growth]</p>

<p><em>Note: This analysis is an interpretation based on limited information and should be taken as exploratory rather than definitive.</em></p>

Important formatting rules:
- Use appropriate HTML tags for headings, paragraphs, and lists as shown.
- Ensure that the final response is valid HTML and can be rendered directly on a web page.
- Do not include any extra text outside the HTML structure.
"""

        try:
            response = psychology_model.generate_content([prompt], safety_settings=safety_settings)
            response_text = response.text.strip()
            return jsonify({'response': response_text})
        except Exception as e:
            logging.error(f"Error generating psychology prediction: {e}")
            return jsonify({'error': "An error occurred while generating the prediction. Please try again."}), 500

    return render_template('psychology_prediction.html')

@app.route('/code_generation', methods=['GET', 'POST'])
def code_generation():
    if request.method == 'POST':
        code_type = request.form['codeType']
        language = request.form['language']
        prompt = f"Write a {language} code to implement {code_type}."
        response = code_model.generate_content([prompt], safety_settings=safety_settings)  # Use code_model here
        if response.candidates and response.candidates[0].content.parts:
            response_text = response.candidates[0].content.parts[0].text
        else:
            response_text = "No valid response found."
        return jsonify({'response': response_text})
    return render_template('code_generation.html')


@app.route('/algorithm_generation', methods=['GET', 'POST'])
def algorithm_generation():
    if request.method == 'POST':
        algo = request.form['algorithm']
        prompt = f"""
        Write a function to implement the {algo} algorithm. Follow these guidelines:
        1. Ensure the function is well-structured and follows best practices for readability and efficiency.
        2. Include clear comments explaining the logic and any complex steps.
        3. Use type hints for function parameters and return values.
        4. Include a brief docstring explaining the purpose of the function and its parameters.
        5. If applicable, include a simple example of how to use the function.
        6. If the algorithm is complex, consider breaking it down into smaller helper functions.
        """
        
        try:
            response = algorithm_model.generate_content([prompt], safety_settings=safety_settings)
            if response.candidates and response.candidates[0].content.parts:
                response_text = response.candidates[0].content.parts[0].text
                # Format the response for better display
                formatted_response = response_text.replace('```python', '<pre><code class="language-python">').replace('```', '</code></pre>')
                return jsonify({'response': formatted_response})
            else:
                return jsonify({'error': "No valid response generated. Please try again."}), 500
        except Exception as e:
            logging.error(f"Error generating algorithm: {e}")
            return jsonify({'error': f"An error occurred: {str(e)}"}), 500

    return render_template('algorithm_generation.html')



from io import BytesIO

@app.route('/analyze', methods=['GET', 'POST'])
def analyze():
    if request.method == 'POST':
        try:
            gender = request.form.get('gender')
            symptoms = request.form.get('symptoms')
            body_part = request.form.get('body-part')
            layer = request.form.get('layer')
            image = request.files.get('image')

            prompt = f"""As an AI medical assistant, analyze the following information about a patient:

Gender: {gender}
Symptoms: {symptoms}
Affected Body Part: {body_part}
Layer Affected: {layer}

Based on this information, provide a detailed analysis considering the following:

1. Possible conditions: List and briefly describe potential conditions that match the symptoms and affected area.
2. Risk factors: Discuss any risk factors associated with the gender or affected body part.
3. Recommended next steps: Suggest appropriate medical tests or examinations that could help diagnose the condition.
4. General advice: Offer some general health advice related to the symptoms or affected area.

Important: This is not a diagnosis. Advise the patient to consult with a healthcare professional for an accurate diagnosis and treatment plan.

Format the response using the following structure:
<section>
<h2>Section Title</h2>
<p>Paragraph text</p>
<ul>
<li>List item 1</li>
<li>List item 2</li>
</ul>
</section>

Use <strong> for emphasis on important points.
"""

            if image:
                img = Image.open(BytesIO(image.read()))
                img_byte_arr = BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_byte_arr = img_byte_arr.getvalue()
                image_base64 = base64.b64encode(img_byte_arr).decode('utf-8')

                prompt += f"""
<section>
<h2>Image Analysis</h2>
<p>Analyze the provided image in relation to the patient's symptoms and affected body part. Consider:</p>
<ul>
<li>Any visible symptoms or abnormalities</li>
<li>Correlation between the image and the reported symptoms</li>
<li>Additional insights the image might provide about the patient's condition</li>
</ul>
</section>

Image data: data:image/png;base64,{image_base64}
"""

                response = model_vision.generate_content([prompt, Image.open(BytesIO(base64.b64decode(image_base64)))], safety_settings=safety_settings)
            else:
                response = model_text.generate_content([prompt], safety_settings=safety_settings)

            analysis_text = response.text if hasattr(response, 'text') else response.parts[0].text

            # Wrap the entire response in a div for styling
            formatted_analysis = f'<div class="analysis-content">{analysis_text}</div>'

            return jsonify({'analysis': formatted_analysis})
        except Exception as e:
            logging.error(f"Error in /analyze route: {e}")
            return jsonify({'error': "Internal Server Error"}), 500
    return render_template('analyze.html')


# Flowchart Generation Routes
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():  # Only include non-empty paragraphs
            full_text.append(para.text)
    return '\n'.join(full_text)

def extract_text_from_pdf(file_path):
    with open(file_path, 'rb') as f:
        pdfReader = PyPDF2.PdfReader(f)
        full_text = []
        for page in pdfReader.pages:
            text = page.extract_text()
            if text.strip():  # Only include non-empty pages
                full_text.append(text)
    return '\n'.join(full_text)

def clean_and_validate_json(text):
    """Clean and validate JSON from the model's response."""
    json_match = re.search(r'\{[\s\S]*\}', text)
    if not json_match:
        return None

    json_str = json_match.group()

    json_str = re.sub(r'```json\s*', '', json_str)
    json_str = re.sub(r'```\s*$', '', json_str)
    json_str = json_str.strip()

    try:
        json_data = json.loads(json_str)

        if not all(key in json_data for key in ['nodes', 'edges']):
            return None

        for node in json_data['nodes']:
            if not all(key in node for key in ['id', 'label']):
                return None
            node['shape'] = node.get('shape', 'box')
            node['level'] = node.get('level', 0)
            node['order'] = node.get('order', 1)

        for edge in json_data['edges']:
            if not all(key in edge for key in ['from', 'to']):
                return None
            edge['order'] = edge.get('order', 1)

        return json_data
    except json.JSONDecodeError:
        return None

def generate_flowchart(topic, chart_type, animation, detail_level, document_text=None):
    max_text_length = 4000
    if document_text:
        topic_prompt = f"Generate a hierarchical {'mind map' if chart_type == 'mind_map' else 'flowchart'} based on this content:\n\n{document_text}\n\n"
    else:
        topic_prompt = f"Generate a hierarchical {'mind map' if chart_type == 'mind_map' else 'flowchart'} for: \"{topic}\".\n\n"

    prompt = topic_prompt + f"""
Please create a {'mind map' if chart_type == 'mind_map' else 'flowchart'} that is {'animated' if animation == 'animated' else 'static'} and {'simple' if detail_level == 'simple' else 'normal' if detail_level == 'normal' else 'detailed'}.

Output a JSON object with this exact structure:
{{
    "nodes": [
        {{"id": 1, "label": "Start", "shape": "ellipse", "level": 0, "order": 1}},
        {{"id": 2, "label": "Process", "shape": "box", "level": 1, "order": 2}}
    ],
    "edges": [
        {{"from": 1, "to": 2, "order": 1}}
    ]
}}

Rules:
1. Use only these shapes: "ellipse", "box", "diamond", "hexagon", "circle"
2. Each node must have a unique integer id
3. Level 0 is root, increasing for each sub-level
4. Include order for animation sequence
5. Keep labels clear and concise
6. Maximum 20 nodes for simple, 35 for normal, 50 for detailed
7. Output ONLY the JSON, no other text"""

    try:
        response = model.generate_content(prompt)
        flowchart_data = clean_and_validate_json(response.text)

        if flowchart_data is None:
            return {"error": "Invalid JSON structure", "raw_response": response.text}

        return flowchart_data
    except Exception as e:
        return {"error": f"Error generating flowchart: {str(e)}"}

def modify_flowchart(current_data, prompt, chart_type):
    """Modifies the current flowchart based on a user prompt."""
    current_data_str = json.dumps(current_data)
    prompt_text = f"""Given the current {'mind map' if chart_type == 'mind_map' else 'flowchart'} data:\n\n{current_data_str}\n\nModify it according to the following prompt: \"{prompt}\".

The output should be a JSON object with the same structure as before, representing the updated {'mind map' if chart_type == 'mind_map' else 'flowchart'}. Ensure that the node and edge IDs remain unique and consistent where applicable.

Output ONLY the JSON, no other text."""

    try:
        response = model.generate_content(prompt_text)
        modified_data = clean_and_validate_json(response.text)

        if modified_data is None:
            return {"error": "Invalid JSON structure from modification", "raw_response": response.text}

        return modified_data
    except Exception as e:
        return {"error": f"Error modifying flowchart: {str(e)}"}

@app.route('/get_flowchart_data', methods=['POST'])
def get_flowchart_data():
    global current_flowchart_data
    try:
        data = request.form
        topic = data.get('topic', '').strip()
        chart_type = data.get('type', 'flowchart')
        animation = data.get('animation', 'static')
        detail_level = data.get('detail_level', 'normal')

        document_text = None
        file = request.files.get('file')

        if file and file.filename:
            if not allowed_file(file.filename):
                return jsonify({"error": "Unsupported file type."}), 400

            filename = secure_filename(file.filename)
            temp_fd, temp_path = tempfile.mkstemp()

            try:
                with os.fdopen(temp_fd, 'wb') as temp_file:
                    file.save(temp_file)

                if filename.lower().endswith('.docx'):
                    document_text = extract_text_from_docx(temp_path)
                elif filename.lower().endswith('.pdf'):
                    document_text = extract_text_from_pdf(temp_path)

                if not topic and document_text:
                    topic = "Flowchart from Document"

            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

        if not topic and not file:
            return jsonify({"error": "Please provide a topic or upload a document."}), 400

        flowchart_data = generate_flowchart(topic, chart_type, animation, detail_level, document_text)

        if 'error' in flowchart_data:
            return jsonify(flowchart_data), 500

        current_flowchart_data = flowchart_data # Store the generated data

        nodes = [{
            "id": node["id"],
            "label": node["label"],
            "shape": node.get("shape", "box"),
            "order": node.get("order", 1),
            "level": node.get("level", 0)
        } for node in flowchart_data.get('nodes', [])]

        edges = [{
            "from": edge["from"],
            "to": edge["to"],
            "id": f"{edge['from']}-{edge['to']}",
            "order": edge.get("order", 1)
        } for edge in flowchart_data.get('edges', [])]

        nodes.sort(key=lambda x: x['order'])
        edges.sort(key=lambda x: x['order'])

        return jsonify({
            "nodes": nodes,
            "edges": edges,
            "animation": animation,
            "chart_type": chart_type
        })

    except Exception as e:
        logging.error(f"Error in get_flowchart_data: {str(e)}")
        return jsonify({"error": "An unexpected error occurred."}), 500

@app.route('/add_node', methods=['POST'])
def add_node():
    global current_flowchart_data
    data = request.get_json()
    new_node = data.get('node')
    if new_node:
        # Simple way to generate a new unique ID (can be improved)
        new_id = max([node['id'] for node in current_flowchart_data['nodes']] or [0]) + 1
        new_node['id'] = new_id
        current_flowchart_data['nodes'].append(new_node)
        return jsonify({"status": "success", "node": new_node})
    return jsonify({"status": "error", "message": "Invalid node data"}), 400

@app.route('/delete_node/<int:node_id>', methods=['DELETE'])
def delete_node(node_id):
    global current_flowchart_data
    current_flowchart_data['nodes'] = [node for node in current_flowchart_data['nodes'] if node['id'] != node_id]
    current_flowchart_data['edges'] = [edge for edge in current_flowchart_data['edges']
                                       if edge['from'] != node_id and edge['to'] != node_id]
    return jsonify({"status": "success"})

@app.route('/edit_node/<int:node_id>', methods=['PUT'])
def edit_node(node_id):
     global current_flowchart_data
     data = request.get_json()
     new_label = data.get('node').get('label')
     for node in current_flowchart_data['nodes']:
          if node['id'] == node_id:
                node['label'] = new_label
                return jsonify({"status": "success", "node": node})
     return jsonify({"status": "error", "message": "Node not found"}), 404

@app.route('/add_edge', methods=['POST'])
def add_edge():
    global current_flowchart_data
    data = request.get_json()
    new_edge = data.get('edge')
    if new_edge:
        current_flowchart_data['edges'].append(new_edge)
        return jsonify({"status": "success", "edge": new_edge})
    return jsonify({"status": "error", "message": "Invalid edge data"}), 400

@app.route('/delete_edge/<from_id>/<to_id>', methods=['DELETE'])
def delete_edge(from_id, to_id):
    global current_flowchart_data
    current_flowchart_data['edges'] = [
        edge for edge in current_flowchart_data['edges']
        if not (str(edge['from']) == from_id and str(edge['to']) == to_id)
    ]
    return jsonify({"status": "success"})

@app.route('/modify_flowchart_prompt', methods=['POST'])
def modify_flowchart_prompt():
    global current_flowchart_data, is_chart_modifying
    data = request.get_json()
    prompt = data.get('prompt')
    chart_type = data.get('chart_type', 'flowchart')

    if not prompt:
        return jsonify({"status": "error", "message": "Prompt cannot be empty"}), 400

    if is_chart_modifying:
      return jsonify({"status": "error", "message": "Chart is currently being modified, please wait..."}), 400

    is_chart_modifying = True  # set flag
    try:
        modified_data = modify_flowchart(current_flowchart_data, prompt, chart_type)

        if 'error' in modified_data:
            return jsonify(modified_data), 500

        current_flowchart_data = modified_data # Update the current data
        
        # Prepare the data for vis-network
        nodes = [{
            "id": node["id"],
            "label": node["label"],
            "shape": node.get("shape", "box"),
            "order": node.get("order", 1),
            "level": node.get("level", 0)
        } for node in modified_data.get('nodes', [])]

        edges = [{
            "from": edge["from"],
            "to": edge["to"],
            "id": f"{edge['from']}-{edge['to']}",
            "order": edge.get("order", 1)
        } for edge in modified_data.get('edges', [])]

        return jsonify({
            "status": "success",
            "nodes": nodes,
            "edges": edges
        })
    except Exception as e:
       return jsonify({"error": f"Error modifying flowchart: {str(e)}"})
    finally:
      is_chart_modifying = False # clear flag


@app.route('/send-email', methods=['POST'])
def send_email():
    data = request.json
    name = data.get('name')
    email = data.get('email')
    message = data.get('message')

    mail_data = {
        'Messages': [
            {
                "From": {
                    "Email": "21131A05C6@gvpce.ac.in",  # Replace with your email
                    "Name": "Kv Nexus"
                },
                "To": [
                    {
                        "Email": "21131A05C6@gvpce.ac.in",  # Replace with your email
                        "Name": "Kv Nexus"
                    }
                ],
                "Subject": f"New Contact Form Submission from {name}",
                "TextPart": f"Name: {name}\nEmail: {email}\nMessage: {message}",
                "HTMLPart": f"<h3>New Contact Form Submission</h3><p><strong>Name:</strong> {name}</p><p><strong>Email:</strong> {email}</p><p><strong>Message:</strong> {message}</p>"
            }
        ]
    }

    result = mailjet.send.create(data=mail_data)
    
    if result.status_code == 200:
        return jsonify({"message": "Email sent successfully!"}), 200
    else:
        return jsonify({"message": "Failed to send email."}), 500


#docuement summarize 

# Token bucket for rate limiting
# Rate limiting parameters
REQUEST_LIMIT = 15
TIME_WINDOW = 60

class TokenBucket:
    def __init__(self, tokens, fill_rate):
        self.capacity = tokens
        self.tokens = tokens
        self.fill_rate = fill_rate
        self.last_check = time.time()
        self.lock = threading.Lock()

    def get_token(self):
        with self.lock:
            now = time.time()
            time_passed = now - self.last_check
            self.tokens = min(self.capacity, self.tokens + time_passed * self.fill_rate)
            self.last_check = now
            if self.tokens >= 1:
                self.tokens -= 1
                return True
            return False

# Initialize the token bucket (15 tokens, refill 1 token every 4 seconds)
token_bucket = TokenBucket(REQUEST_LIMIT, 1 / (TIME_WINDOW / REQUEST_LIMIT))

def rate_limit_check():
    while not token_bucket.get_token():
        time.sleep(1)
# Rate limiting parameters

rate_limit_lock = None  # Replace with your lock mechanism
last_reset_time = time.time()
request_count = 0

def rate_limited(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        rate_limit_check()
        return func(*args, **kwargs)
    return wrapper


import fitz  # PyMuPDF
import io
from PIL import Image
import base64

def process_page(pdf_document, page_num, doc_ref):
    logger.info(f"Processing page {page_num + 1}")
    page = pdf_document[page_num]
    
    # Convert page to image
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Increase resolution
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    
    # Convert image to base64
    buffered = io.BytesIO()
    img.save(buffered, format="PNG")
    img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')

    try:
        page_summary = generate_summary(img_base64)
        if page_summary:
            logger.info(f"Summary generated for page {page_num + 1}")
            doc_ref.update({
                'current_page': page_num + 1,
                'summary': firestore.ArrayUnion([page_summary])
            })
        else:
            logger.warning(f"Failed to generate summary for page {page_num + 1}")
            doc_ref.update({
                'current_page': page_num + 1,
                'summary': firestore.ArrayUnion([f"(Summary not available for page {page_num + 1})"])
            })
    except Exception as e:
        logger.error(f"Error processing page {page_num + 1}: {e}")
        doc_ref.update({
            'current_page': page_num + 1,
            'summary': firestore.ArrayUnion([f"(Error processing page {page_num + 1})"])
        })

@app.route('/document_summarizer', methods=['GET', 'POST'])
def document_summarizer():
    return render_template('document_summarizer.html')

@app.route('/quote', methods=['GET'])
def get_quote():
    import random
    quotes = [
        "The best way to predict the future is to invent it. – Alan Kay",
        "Life is like riding a bicycle. To keep your balance you must keep moving. – Albert Einstein",
        "Problems are not stop signs, they are guidelines. – Robert H. Schuller",
        "In order to succeed, we must first believe that we can. – Nikos Kazantzakis",
        "The only limit to our realization of tomorrow is our doubts of today. – Franklin D. Roosevelt"
    ]
    quote = random.choice(quotes)
    return jsonify({'quote': quote})

@rate_limited
def generate_summary(image_base64):
    rate_limit_check()  # Wait for a token before making the API call
    
    prompt = [
        """Analyze the following image, which is a page from a document, and provide a concise and simplified summary. Ensure the summary is well-structured with clear headings and subheadings.

Formatting Guidelines:

- Use `#` for main section titles.
- Use `##` for subsections.
- Use `-` for bullet points.
- For **bold text**, wrap the text with double asterisks, e.g., `**important**`.
- For *italic text*, wrap the text with single asterisks, e.g., `*note*`.
- **For tables**, use proper Markdown table syntax with pipes `|` and hyphens `-` for headers.

- Keep sentences short and use simple language.
- Focus on the main ideas and avoid unnecessary details.
- Do not include direct error messages or irrelevant information.

Here is the image to analyze and summarize:
""",
        Image.open(io.BytesIO(base64.b64decode(image_base64)))
    ]

    try:
        response = model_vision.generate_content(prompt, safety_settings=safety_settings)
        summary_text = response.text
        logger.info("Summary generated successfully")
        return summary_text
    except google.api_core.exceptions.ResourceExhausted as e:
        logger.warning(f"Resource exhausted: {e}. Retrying...")
        raise  # This will trigger a retry
    except Exception as e:
        logger.error(f"Error in Gemini API call: {e}")
        return None  # Return None for non-retryable errors

def create_word_document(summary):
    doc = Document()

    # Define styles
    define_custom_styles(doc)

    # Adjust document layout
    adjust_document_layout(doc)

    # Set page borders
    set_page_border(doc)

    # Convert Markdown to HTML with necessary extensions
    html = markdown.markdown(summary, extensions=['extra', 'tables', 'fenced_code', 'codehilite', 'nl2br'])

    # Parse HTML
    soup = BeautifulSoup(html, 'html.parser')

    # Iterate over HTML elements and add them to the Word document
    for element in soup.contents:
        process_html_element(doc, element)

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def define_custom_styles(doc):
    styles = doc.styles

    # Title Style (Heading 1)
    style_h1 = styles['Heading 1']
    style_h1.font.name = 'Calibri Light'
    style_h1.font.size = Pt(24)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(31, 56, 100)
    style_h1.paragraph_format.space_after = Pt(12)
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle Style (Heading 2)
    style_h2 = styles['Heading 2']
    style_h2.font.name = 'Calibri'
    style_h2.font.size = Pt(20)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(54, 95, 145)
    style_h2.paragraph_format.space_before = Pt(12)
    style_h2.paragraph_format.space_after = Pt(6)

    # Normal Text Style
    style_normal = styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.space_after = Pt(8)
    style_normal.paragraph_format.line_spacing = 1.15

    # List Bullet Style
    style_list_bullet = styles['List Bullet']
    style_list_bullet.font.name = 'Calibri'
    style_list_bullet.font.size = Pt(12)
    style_list_bullet.paragraph_format.space_after = Pt(4)
    style_list_bullet.paragraph_format.line_spacing = 1.15

    # List Number Style
    style_list_number = styles['List Number']
    style_list_number.font.name = 'Calibri'
    style_list_number.font.size = Pt(12)
    style_list_number.paragraph_format.space_after = Pt(4)
    style_list_number.paragraph_format.line_spacing = 1.15

def adjust_document_layout(doc):
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

def set_page_border(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_position in ('top', 'left', 'bottom', 'right'):
            border_el = OxmlElement(f'w:{border_position}')
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '24')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), '5B9BD5')
            pgBorders.append(border_el)
        sectPr.append(pgBorders)

def process_html_element(doc, element, parent=None):
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip():
            if parent is not None:
                parent.add_run(text)
            else:
                p = doc.add_paragraph()
                p.add_run(text)
        return
    elif element.name is not None:
        if element.name == 'h1':
            p = doc.add_heading(level=1)
            add_runs_from_element(p, element)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.name == 'h2':
            p = doc.add_heading(level=2)
            add_runs_from_element(p, element)
        elif element.name == 'p':
            p = doc.add_paragraph()
            add_runs_from_element(p, element)
        elif element.name in ['strong', 'b']:
            if parent is not None:
                run = parent.add_run(element.get_text())
                run.bold = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.bold = True
        elif element.name in ['em', 'i']:
            if parent is not None:
                run = parent.add_run(element.get_text())
                run.italic = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.italic = True
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                add_runs_from_element(p, li)
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Number')
                add_runs_from_element(p, li)
        elif element.name == 'table':
            add_table_to_document_from_html(doc, element)
        else:
            # Process children
            for child in element.contents:
                process_html_element(doc, child, parent)
    else:
        # Unknown element type
        pass

def add_runs_from_element(paragraph, element):
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip():
            paragraph.add_run(text)
    elif element.name is not None:
        if element.name in ['strong', 'b']:
            run = paragraph.add_run(element.get_text())
            run.bold = True
        elif element.name in ['em', 'i']:
            run = paragraph.add_run(element.get_text())
            run.italic = True
        else:
            for content in element.contents:
                add_runs_from_element(paragraph, content)
    else:
        # Unknown element type
        pass

def add_table_to_document_from_html(doc, table_element):
    rows = table_element.find_all('tr')
    if not rows:
        return

    num_cols = len(rows[0].find_all(['td', 'th']))
    table = doc.add_table(rows=0, cols=num_cols)
    table.style = 'Table Grid'  # Or define a custom table style

    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        row_cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            cell_text = cell.get_text(strip=True)
            paragraph = row_cells[idx].paragraphs[0]
            paragraph.clear()  # Clear existing content
            run = paragraph.add_run(cell_text)
            row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Apply styling for header cells
            if cell.name == 'th':
                run.bold = True
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'D9E1F2')  # Light blue background
                row_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)

@app.route('/upload', methods=['POST'])
@rate_limited
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file and file.filename.lower().endswith('.pdf'):
        try:
            # Read the file into memory
            file_content = file.read()
            file_size = len(file_content)

            # Check file size (10MB limit)
            if file_size > 10 * 1024 * 1024:
                return jsonify({'error': 'File size exceeds 10MB limit'}), 400

            # Generate a unique PDF ID
            pdf_id = str(uuid.uuid4())

            # Upload the PDF to Firebase Storage
            blob = bucket.blob(f'pdfs/{pdf_id}.pdf')
            blob.upload_from_string(file_content, content_type='application/pdf')

            # Get the total number of pages
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            total_pages = len(pdf_document)
            pdf_document.close()

            # Initialize processing status in Firestore
            db.collection('pdf_processes').document(pdf_id).set({
                'status': 'processing',
                'current_page': 0,
                'total_pages': total_pages,
                'summary': '',
                'processing_start_time': time.time(),
                'timestamp': firestore.SERVER_TIMESTAMP,
                'file_size': file_size
            })

            return jsonify({
                'pdf_id': pdf_id,
                'total_pages': total_pages,
                'file_size': file_size
            }), 200
        except Exception as e:
            logging.error(f"Error uploading file: {e}")
            return jsonify({'error': f'Error uploading file: {str(e)}'}), 500
    else:
        return jsonify({'error': 'Invalid file type. Please upload a PDF.'}), 400

# Update the process_pdf_endpoint function
@app.route('/process_pdf', methods=['POST'])
def process_pdf_endpoint():
    data = request.get_json()
    pdf_id = data.get('pdf_id')
    if not pdf_id:
        logger.error("No PDF ID provided")
        return jsonify({'error': 'No PDF ID provided.'}), 400

    doc_ref = db.collection('pdf_processes').document(pdf_id)
    doc = doc_ref.get()
    if not doc.exists:
        logger.error(f"Invalid PDF ID: {pdf_id}")
        return jsonify({'error': 'Invalid PDF ID.'}), 400
    
    result = doc.to_dict()
    current_page = result['current_page']
    total_pages = result['total_pages']

    if current_page >= total_pages:
        logger.info(f"PDF {pdf_id} processing already completed")
        return jsonify({'status': 'completed'}), 200

    try:
        logger.info(f"Processing PDF {pdf_id}, page {current_page + 1} of {total_pages}")
        blob = bucket.blob(f'pdfs/{pdf_id}.pdf')
        pdf_bytes = blob.download_as_bytes()
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")

        process_page(pdf_document, current_page, doc_ref)
        pdf_document.close()

        updated_doc = doc_ref.get().to_dict()
        if updated_doc['current_page'] >= total_pages:
            logger.info(f"PDF {pdf_id} processing completed")
            doc_ref.update({
                'status': 'completed',
                'processing_end_time': time.time()
            })
            blob.delete()
            return jsonify({'status': 'completed'}), 200
        else:
            logger.info(f"PDF {pdf_id} processing in progress. Current page: {updated_doc['current_page']}")
            return jsonify({
                'status': 'processing',
                'current_page': updated_doc['current_page'],
                'total_pages': total_pages
            }), 200

    except Exception as e:
        logger.error(f"Error processing PDF {pdf_id}: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/check_status', methods=['GET'])
def check_status():
    pdf_id = request.args.get('pdf_id')
    if not pdf_id:
        logger.error("No PDF ID provided for status check")
        return jsonify({'error': 'No PDF ID provided.'}), 400

    doc_ref = db.collection('pdf_processes').document(pdf_id)
    doc = doc_ref.get()
    if not doc.exists:
        logger.error(f"Invalid PDF ID for status check: {pdf_id}")
        return jsonify({'error': 'Invalid PDF ID.'}), 400
    
    result = doc.to_dict()
    status = result.get('status', 'processing')

    if status == 'completed':
        logger.info(f"Status check: PDF {pdf_id} processing completed")
        summary = '\n\n'.join(result['summary'])
        docx_buffer = create_word_document(summary)
        docx_base64 = base64.b64encode(docx_buffer.getvalue()).decode('utf-8')

        processing_time = 'N/A'
        if result.get('processing_start_time') and result.get('processing_end_time'):
            processing_time = int(result['processing_end_time'] - result['processing_start_time'])

        return jsonify({
            'status': 'completed',
            'docx': docx_base64,
            'total_pages': result.get('total_pages', 0),
            'processing_time': processing_time
        }), 200
    else:
        logger.info(f"Status check: PDF {pdf_id} processing in progress. Current page: {result.get('current_page', 0)}")
        return jsonify({
            'status': 'processing',
            'current_page': result.get('current_page', 0),
            'total_pages': result.get('total_pages', 0)
        }), 200

# --- Story Generation and Editing Functions ---
def generate_story(keywords, genre, num_chapters, tone, style, age_group, include_magic, include_superpowers, include_conflict, image_style):
    """Generate story content using Gemini API with enhanced parameters."""
    prompt = f"""
    Create a captivating story based on the following:
    Keywords: {keywords}
    Genre: {genre}
    Number of chapters: {num_chapters}
    Tone: {tone}
    Writing Style: {style}
    Target Audience Age Group: {age_group}
    Include Magical Elements: {'Yes' if include_magic else 'No'}
    Include Superpowers/Special Abilities: {'Yes' if include_superpowers else 'No'}
    Include Major Conflicts: {'Yes' if include_conflict else 'No'}
    
    Format the response exactly like this JSON structure (don't include any other text before or after the JSON):
    {{
        "title": "Story title",
        "author": "AI Author",
        "moral": "The moral of the story",
        "chapters": [
            {{
                "chapter_number": 1,
                "chapter_title": "Chapter title",
                "content": "Chapter content (approximately 2-4 paragraphs, suitable for the target age group and writing style)",
                "image_prompt": "Detailed visual description for a captivating chapter image, reflecting the chapter's mood and style",
                "terminology": {{}}
            }}
        ]
    }}
    """

    try:
        response = final_story_generation_model.generate_content(prompt)
        response_text = response.text

        # Parse JSON with multiple fallback attempts
        try:
            story_data = json.loads(response_text)
        except json.JSONDecodeError:
            # Try finding JSON in markdown code blocks
            json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
            if json_match:
                story_data = json.loads(json_match.group(1))
            else:
                # Try finding content between curly braces
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    story_data = json.loads(json_match.group(0))
                else:
                    story_data = {
                        "title": f"{genre.title()} Story",
                        "author": "AI Author",
                        "moral": "Could not generate a moral for the story",
                        "chapters": [{
                            "chapter_number": 1,
                            "chapter_title": "Chapter 1",
                            "content": "Story generation failed. Please try again.",
                            "image_prompt": "A blank page with some text",
                            "terminology": {}
                        }]
                    }

        # Add terminology to each chapter
        for chapter in story_data['chapters']:
            chapter['terminology'] = extract_terminology(chapter['content'])

        return story_data

    except Exception as e:
        print(f"Error generating story: {str(e)}")
        return {
            "title": f"{genre.title()} Story",
            "author": "AI Author",
            "moral": "Could not generate a moral for the story",
            "chapters": [{
                "chapter_number": 1,
                "chapter_title": "Chapter 1",
                "content": "Story generation failed. Please try again.",
                "image_prompt": "A blank page with some text",
                "terminology": {}
            }]
        }
def regenerate_chapter(story_data, chapter_number, tone, style, include_magic, include_superpowers, include_conflict, image_style):
    """Regenerate a specific chapter."""
    chapter_index = chapter_number - 1
    if not (0 <= chapter_index < len(story_data["chapters"])):
        return "Invalid chapter number", 400
    
    # Previous context for regeneration
    previous_context = ""
    if chapter_index > 0:
        previous_context = story_data['chapters'][chapter_index - 1]['content']

    prompt = f"""
    Regenerate chapter {chapter_number} of the story titled '{story_data['title']}'.
    
    Previous chapter context (if applicable): {previous_context}
    
    Overall Tone: {tone}
    Writing Style: {style}
    Include Magical Elements: {'Yes' if include_magic else 'No'}
    Include Superpowers/Special Abilities: {'Yes' if include_superpowers else 'No'}
    Include Major Conflicts: {'Yes' if include_conflict else 'No'}

    Format the response exactly like this JSON structure (don't include any other text before or after the JSON):
    {{
        "chapter_number": {chapter_number},
        "chapter_title": "New chapter title",
        "content": "New chapter content (approximately 2-4 paragraphs, consistent with the story's style and tone)",
        "image_prompt": "Detailed visual description for a captivating chapter image, furthering the narrative visually",
        "terminology": {{}}
    }}
    """

    try:
        response = final_story_generation_model.generate_content(prompt)
        try:
            new_chapter = json.loads(response.text)
        except json.JSONDecodeError:
            json_match = re.search(r'```json\s*(.*?)\s*```', response.text, re.DOTALL)
            if json_match:
                new_chapter = json.loads(json_match.group(1))
            else:
                json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                if json_match:
                    new_chapter = json.loads(json_match.group(0))
                else:
                    new_chapter = {
                        "chapter_number": chapter_number,
                        "chapter_title": "Regenerated Chapter",
                        "content": "Failed to regenerate chapter. Please try again.",
                        "image_prompt": "A blank page with some text",
                        "terminology": {}
                    }

        # Add terminology
        new_chapter['terminology'] = extract_terminology(new_chapter['content'])

        # Replace old chapter with new chapter
        story_data['chapters'][chapter_index] = new_chapter

        # Regenerate image for the chapter
        new_chapter['image'] = generate_image(new_chapter['image_prompt'], 800, 600, image_style)

        return new_chapter

    except Exception as e:
        print(f"Error regenerating chapter: {str(e)}")
        return {
            "chapter_number": chapter_number,
            "chapter_title": "Regenerated Chapter",
            "content": "Failed to regenerate chapter. Please try again.",
            "image_prompt": "A blank page with some text",
            "terminology": {}
        }
def continue_story(previous_story, num_new_chapters, tone, style, include_magic, include_superpowers, include_conflict, image_style):
    """Generate additional chapters for an existing story with enhanced parameters."""
    prompt = f"""
    Continue this story with {num_new_chapters} more chapters, maintaining the established themes and characters.
    Previous story title: {previous_story['title']}
    Last chapter content: {previous_story['chapters'][-1]['content']}
    Overall Tone: {tone}
    Writing Style: {style}
    Include Magical Elements: {'Yes' if include_magic else 'No'}
    Include Superpowers/Special Abilities: {'Yes' if include_superpowers else 'No'}
    Include Major Conflicts: {'Yes' if include_conflict else 'No'}

    Format the response exactly like this JSON structure (don't include any other text before or after the JSON):
    {{
        "chapters": [
            {{
                "chapter_number": {len(previous_story['chapters']) + 1},
                "chapter_title": "New chapter title",
                "content": "New chapter content (approximately 2-4 paragraphs, consistent with the story's style and tone)",
                "image_prompt": "Detailed visual description for a captivating chapter image, furthering the narrative visually",
                "terminology": {{}}
            }}
        ]
    }}
    """

    try:
        response = final_story_generation_model.generate_content(prompt)
        try:
            new_chapters = json.loads(response.text)
        except json.JSONDecodeError:
            json_match = re.search(r'```json\s*(.*?)\s*```', response.text, re.DOTALL)
            if json_match:
                new_chapters = json.loads(json_match.group(1))
            else:
                json_match = re.search(r'\{.*\}', response.text, re.DOTALL)
                if json_match:
                    new_chapters = json.loads(json_match.group(0))
                else:
                    new_chapters = {
                        "chapters": [{
                            "chapter_number": len(previous_story['chapters']) + 1,
                            "chapter_title": "New Chapter",
                            "content": "Failed to generate new content. Please try again.",
                            "image_prompt": "A blank page with some text",
                            "terminology": {}
                        }]
                    }

        # Add terminology to each new chapter
        for chapter in new_chapters['chapters']:
            chapter['terminology'] = extract_terminology(chapter['content'])

        return new_chapters['chapters']
    except Exception as e:
        print(f"Error continuing story: {str(e)}")
        return [{
            "chapter_number": len(previous_story['chapters']) + 1,
            "chapter_title": "New Chapter",
            "content": "Failed to generate new content. Please try again.",
            "image_prompt": "A blank page with some text",
            "terminology": {}
        }]

def get_word_definitions(words):
    """Get definitions using Gemini API with improved prompting and error handling."""
    prompt = f"""
    Define the following words clearly and concisely, focusing on their most common meanings in everyday usage.
    Words to define: {', '.join(words)}

    Respond ONLY with a JSON object in this exact format (no other text):
    {{
        "word1": "simple definition here",
        "word2": "simple definition here"
    }}

    Make sure each definition is:
    - Clear and simple
    - 10-15 words maximum
    - Suitable for the general audience
    - Focuses on the most common meaning
    """
    
    try:
        response = final_story_generation_model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Clean up the response to ensure it's valid JSON
        # Remove any markdown formatting if present
        if '```json' in response_text:
            response_text = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL).group(1)
        elif '```' in response_text:
            response_text = re.search(r'```\s*(.*?)\s*```', response_text, re.DOTALL).group(1)
            
        # Remove any remaining non-JSON text
        response_text = re.search(r'\{.*\}', response_text, re.DOTALL).group(0)
        
        definitions = json.loads(response_text)
        
        # Ensure all requested words have definitions
        final_definitions = {}
        for word in words:
            word_lower = word.lower()
            # Try to find the word in the definitions (case-insensitive)
            matching_key = next((k for k in definitions.keys() if k.lower() == word_lower), None)
            if matching_key and definitions[matching_key].strip():
                final_definitions[word] = definitions[matching_key]
            else:
                # If definition is missing, make another attempt for just this word
                single_word_prompt = f"""
                Define this word clearly and concisely in 10-15 words:
                Word: {word}

                Respond ONLY with the definition (no other text).
                """
                try:
                    retry_response = final_story_generation_model.generate_content(single_word_prompt)
                    definition = retry_response.text.strip()
                    if definition:
                        final_definitions[word] = definition
                    else:
                        final_definitions[word] = "No definition available"
                except Exception:
                    final_definitions[word] = "No definition available"
                    
        return final_definitions
        
    except Exception as e:
        print(f"Error getting definitions: {str(e)}")
        return {word: f"Definition not available due to error: {str(e)}" for word in words}

def extract_terminology(text):
    """Extract 4-5 significant terms and get their definitions using Gemini."""
    # Find words that are potentially complex or important
    words = re.findall(r'\b[A-Za-z]{5,}\b', text)
    
    # Remove common words and duplicates
    common_words = {
        'there', 'their', 'would', 'could', 'should', 'about', 'which', 'these', 
        'those', 'were', 'have', 'that', 'what', 'when', 'where', 'while', 'from',
        'been', 'being', 'other', 'another', 'every', 'everything', 'something',
        'anything', 'nothing', 'through', 'although', 'though', 'without', 'within',
        'around', 'before', 'after', 'under', 'over', 'because'
    }
    
    # Filter words
    filtered_words = []
    for word in words:
        word_lower = word.lower()
        if (
            word_lower not in common_words and
            not word.isupper() and  # Skip acronyms
            len(word) >= 6 and  # Focus on longer words
            not any(char.isdigit() for char in word)  # Skip words with numbers
        ):
            filtered_words.append(word)
    
    # Remove duplicates while preserving case
    unique_words = []
    seen_words = set()
    for word in filtered_words:
        if word.lower() not in seen_words:
            unique_words.append(word)
            seen_words.add(word.lower())
    
    # Select the most interesting words (prioritize longer, less common words)
    selected_words = sorted(unique_words, key=lambda x: (len(x), x.lower()), reverse=True)[:5]
    
    if selected_words:
        # Get definitions using improved Gemini function
        definitions = get_word_definitions(selected_words)
        return {k: v for k, v in definitions.items() if v and v != "No definition available"}
    return {}

def generate_image(prompt, width=800, height=600, style="photographic"):
    """Generate image using Pollinations AI with image style."""
    try:
        style_map = {
            "photographic": "style photographic",
            "anime": "style anime",
            "cinematic": "style cinematic",
            "watercolor": "style watercolor illustration",
            "pencil sketch": "style pencil sketch",
            "pop art": "style pop art",
            "steampunk": "style steampunk",
            "3d render": "style 3d render",
        }
        style_prompt = style_map.get(style, "style photographic")
        full_prompt = f"{prompt}, {style_prompt}"

        encoded_prompt = quote(full_prompt)
        image_url = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width={width}&height={height}&nologo=true"
        return image_url
    except Exception as e:
        print(f"Image generation error: {str(e)}")
        return "https://image.pollinations.ai/prompt/scenic%20view?width=800&height=600&nologo=true"
def generate_all_images_concurrent(story_data, image_style):
    """Generate all images concurrently using ThreadPoolExecutor."""
    try:
        # Create tasks for cover and chapter images
        cover_prompt = f"A professional book cover with a title on it, '{story_data['title']}', fantasy art"
        cover_future = executor.submit(generate_image, cover_prompt, 400, 550, image_style)

        chapter_futures = [executor.submit(generate_image, chapter['image_prompt'], 800, 600, image_style) for chapter in story_data['chapters']]

        # Get results
        cover_image = cover_future.result()
        chapter_images = [future.result() for future in chapter_futures]

        # Assign URLs to story data
        for chapter, image_url in zip(story_data['chapters'], chapter_images):
            chapter['image'] = image_url

        return story_data, cover_image

    except Exception as e:
        print(f"Error in generate_all_images_concurrent: {str(e)}")
        fallback_url = "https://image.pollinations.ai/prompt/scenic%20view?width=800&height=600&nologo=true"
        cover_fallback = "https://image.pollinations.ai/prompt/book%20cover?width=400&height=550&nologo=true"

        for chapter in story_data['chapters']:
            chapter['image'] = fallback_url

        return story_data, cover_fallback

@app.route('/flowchart')
def flowchart():
    return render_template('flowchart.html')  # Changed to flowchart.html

@app.route('/story_generation')
def story_generation():
    return render_template('story_generation.html', responsive_voice_key=responsive_voice_key)

@app.route('/generate', methods=['POST'])
def generate():
    try:
        data = request.get_json()
        keywords = data.get('keywords')
        genre = data.get('genre')
        story_length = data.get('storyLength')
        tone = data.get('tone')
        style = data.get('style')
        age_group = data.get('ageGroup')
        include_magic = data.get('includeMagic')
        include_superpowers = data.get('includeSuperpowers')
        include_conflict = data.get('includeConflict')
        image_style = data.get('imageStyle')  # Get the image style

        # Determine the number of chapters based on story length
        num_chapters = {
            'short': 3,
            'medium': 5,
            'long': 7,
            'grand': 10
        }.get(story_length, 3)

        # Generate story content
        story = generate_story(keywords, genre, num_chapters, tone, style, age_group, include_magic, include_superpowers, include_conflict, image_style)

        # Generate all images concurrently
        story, cover_image = generate_all_images_concurrent(story, image_style)

        return jsonify({
            'story': story,
            'cover_image': cover_image
        })
    except Exception as e:
        print(f"Error in generate route: {str(e)}")
        return jsonify({
            'error': 'Failed to generate story',
            'details': str(e)
        }), 500

@app.route('/regenerate', methods=['POST'])
def regenerate():
    try:
        data = request.get_json()
        story_data = data.get('story')
        chapter_number = int(data.get('chapter_number'))
        tone = data.get('tone')
        style = data.get('style')
        include_magic = data.get('includeMagic')
        include_superpowers = data.get('includeSuperpowers')
        include_conflict = data.get('includeConflict')
        image_style = data.get('imageStyle')  # Get the image style

        new_chapter = regenerate_chapter(story_data, chapter_number, tone, style, include_magic, include_superpowers, include_conflict, image_style)

        return jsonify({'story': story_data, 'new_chapter': new_chapter})

    except Exception as e:
        print(f"Error in regenerate route: {str(e)}")
        return jsonify({
            'error': 'Failed to regenerate chapter',
            'details': str(e)
        }), 500

@app.route('/continue', methods=['POST'])
def continue_story_route():
    try:
        data = request.get_json()
        previous_story = data.get('previous_story')
        num_new_chapters = int(data.get('num_new_chapters', 3))
        tone = data.get('tone')
        style = data.get('style')
        include_magic = data.get('includeMagic')
        include_superpowers = data.get('includeSuperpowers')
        include_conflict = data.get('includeConflict')
        image_style = data.get('imageStyle')  # Get the image style

        # Generate new chapters
        new_chapters = continue_story(previous_story, num_new_chapters, tone, style, include_magic, include_superpowers, include_conflict, image_style)

        # Generate images for new chapters concurrently
        futures = [executor.submit(generate_image, chapter['image_prompt'], 800, 600, image_style)
                  for chapter in new_chapters]

        # Wait for all image generations to complete
        for chapter, future in zip(new_chapters, futures):
            chapter['image'] = future.result()

        return jsonify({'new_chapters': new_chapters})
    except Exception as e:
        print(f"Error in continue route: {str(e)}")
        return jsonify({
            'error': 'Failed to continue story',
            'details': str(e)
        }), 500
@app.route('/get_moral', methods=['POST'])
def get_moral():
    try:
        data = request.get_json()
        story_data = data.get('story')
        return jsonify({'moral': story_data['moral']})
    except Exception as e:
        print(f"Error in get_moral route: {str(e)}")
        return jsonify({
            'error': 'Failed to retrieve moral',
            'details': str(e)
        }), 500

@app.route('/download', methods=['POST'])
def download_pdf():
    try:
        data = request.get_json()
        story_data = data.get('story')
        cover_image = data.get('cover_image')
        
        # Configure PDF options
        pdf_options = {
            'page-size': 'A4',
            'margin-top': '2.5cm',
            'margin-right': '2cm',
            'margin-bottom': '2.5cm',
            'margin-left': '2cm',
            'encoding': 'UTF-8',
        }
        
        # Generate HTML
        pdf_html = render_template('pdf_template.html', story=story_data, cover_image=cover_image)
        
        # Create PDF
        pdf_buffer = BytesIO()
        pisa.CreatePDF(
            pdf_html, 
            dest=pdf_buffer,
            encoding='utf-8',
        )
        
        pdf_buffer.seek(0)
        
        # Create response
        response = make_response(pdf_buffer.read())
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename="{story_data["title"].replace(" ", "_")}.pdf"'
        
        return response
        
    except Exception as e:
        print(f"Download PDF Error: {str(e)}")
        return jsonify({
            "error": "Error generating PDF",
            "details": str(e)
        }), 500

if __name__ == '__main__':
    app.run(debug=True)
